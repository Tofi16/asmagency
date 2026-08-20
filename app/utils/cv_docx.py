"""
Builds a Word (.docx) version of an employer-facing CV — a plainer, editable
counterpart to the styled admin/cv_print.html page (which stays the way to
get a PDF, via the browser's own Print dialog).

Kept deliberately simple/portable: one bilingual (EN label / AR label) info
table per section, in the same order as the printed CV, so anyone opening it
in Word/LibreOffice/Google Docs sees a familiar, professional layout they can
still tweak by hand before sending it to a partner agency.

The 3 CV photos (Small ID, Full-Length, Passport) are embedded at the same
positions they appear on the printed CV: portrait near the top, full-length
photo next to Personal Information, passport scan next to Passport Details.
Only JPG/PNG can be embedded as an inline image — if a slot's file is a PDF
(passport scans are sometimes uploaded as PDF), that one photo is skipped
here with a text note, since Word can't inline a PDF page as a picture; the
original file is still on record and shown on the print/PDF version.
"""
import io
import os

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.image.image import Image as DocxImage

from app.models.cv_profile import CVProfile

AUTHORITY = RGBColor(0x1A, 0x1A, 0x2E)
PRO_BLUE = RGBColor(0x25, 0xAA, 0xE2)
INK_MUTED = RGBColor(0x5B, 0x61, 0x72)
LINE_GREY = "E2E6ED"
CANVAS_GREY = "F5F7FA"


def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text if text not in (None, "") else "—")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def _section_title(doc, en, ar):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1A1A2E")
    p._p.get_or_add_pPr().append(shd)
    run_en = p.add_run(f"  {en.upper()}")
    run_en.font.size = Pt(11)
    run_en.font.bold = True
    run_en.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_ar = p.add_run(f"   {ar}  ")
    run_ar.font.size = Pt(10)
    run_ar.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEE)
    return p


def _info_table(doc, rows):
    """rows: list of (label_en, label_ar, value) -> a 3-col bordered table."""
    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(4.2), Cm(2.6), Cm(9.0))
    for r, (label_en, label_ar, value) in enumerate(rows):
        row = table.rows[r]
        for c, w in enumerate(widths):
            row.cells[c].width = w
        _set_cell_text(row.cells[0], label_en, bold=True, size=9.5)
        _shade_cell(row.cells[0], CANVAS_GREY)
        _set_cell_text(row.cells[1], label_ar, size=9, color=INK_MUTED, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _shade_cell(row.cells[1], CANVAS_GREY)
        _set_cell_text(row.cells[2], str(value) if value not in (None, "") else "—", size=9.5)
    _set_table_borders(table)
    return table


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), LINE_GREY)
        borders.append(el)
    tblPr.append(borders)


def _rating_row(doc, label, level, pct_map):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{label}:  ")
    run.font.bold = True
    run.font.size = Pt(9.5)
    pct = pct_map.get(level, 0)
    bar_len = max(1, round(pct / 5))
    bar_run = p.add_run("█" * bar_len + "░" * (20 - bar_len))
    bar_run.font.size = Pt(9)
    bar_run.font.color.rgb = PRO_BLUE
    level_run = p.add_run(f"  {level}")
    level_run.font.size = Pt(9)
    level_run.font.italic = True
    level_run.font.color.rgb = INK_MUTED


def _fitted_size_cm(file_path, max_width_cm, max_height_cm):
    """Reads the image's native pixel aspect ratio and returns (width_cm,
    height_cm) scaled to fit entirely within the given box — the whole photo
    always ends up on the page, never cropped, whatever its own proportions."""
    try:
        img = DocxImage.from_file(file_path)
        native_w, native_h = img.px_width, img.px_height
    except Exception:
        return max_width_cm, max_height_cm  # unreadable dimensions — fall back to the box itself

    ratio = min(max_width_cm / native_w, max_height_cm / native_h)
    return native_w * ratio, native_h * ratio


def _add_photo_centered(doc, file_path, width_cm, caption=None):
    """Adds one centered image, or a small text note if there's no usable
    (JPG/PNG) file for that slot. Returns True if an actual image was added."""
    ext = file_path.rsplit(".", 1)[-1].lower() if file_path and "." in file_path else ""
    has_image = bool(file_path) and os.path.exists(file_path) and ext in ("jpg", "jpeg", "png")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if has_image:
        run = p.add_run()
        # Square box for the small ID photo — fit-within, so nothing is ever cropped.
        w_cm, h_cm = _fitted_size_cm(file_path, width_cm, width_cm)
        run.add_picture(file_path, width=Cm(w_cm), height=Cm(h_cm))
    else:
        note = "photo not on file" if not file_path else "on file as PDF — see Verify Documents"
        run = p.add_run(f"[{note}]")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = INK_MUTED

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cap_run = cap.add_run(caption)
        cap_run.font.size = Pt(8)
        cap_run.font.bold = True
        cap_run.font.color.rgb = AUTHORITY
    return has_image


def _add_photo_pair(doc, left, right, box_width_cm=8.0, box_height_cm=9.5):
    """left/right: (file_path, caption) tuples. Places both photos side by
    side in a borderless 2-col table, each scaled to fit entirely within its
    own box (never cropped) and centered in its cell — mirrors the print
    page's large photo-showcase row."""
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, (file_path, caption) in enumerate([left, right]):
        cell = table.rows[0].cells[col]
        cell.width = Cm(box_width_cm + 0.5)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ext = file_path.rsplit(".", 1)[-1].lower() if file_path and "." in file_path else ""
        if file_path and os.path.exists(file_path) and ext in ("jpg", "jpeg", "png"):
            run = para.add_run()
            w_cm, h_cm = _fitted_size_cm(file_path, box_width_cm, box_height_cm)
            run.add_picture(file_path, width=Cm(w_cm), height=Cm(h_cm))
        else:
            note = "not on file" if not file_path else "on file as PDF — see Verify Documents"
            run = para.add_run(f"[{note}]")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = INK_MUTED

        cap_cell = table.rows[1].cells[col]
        cap_para = cap_cell.paragraphs[0]
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(8.5)
        cap_run.font.bold = True
        cap_run.font.color.rgb = AUTHORITY
    return table


def build_cv_docx(applicant, profile, company_name, company_phone, company_email, company_address_en,
                   portrait_doc=None, full_doc=None, passport_doc=None):
    """Returns a BytesIO buffer containing the finished .docx — ready to send with Flask's send_file().

    portrait_doc / full_doc / passport_doc are the applicant's Document rows
    for the 3 CV photo slots (or None if that slot hasn't been filled in) —
    pass whatever admin/routes.py::cv_export_docx looked up, same as it does
    for the print page."""
    doc = DocxDocument()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # --- header ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("JOB APPLICATION — CV")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = AUTHORITY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("طلب توظيف — السيرة الذاتية")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = INK_MUTED

    # --- small ID/portrait photo, right under the title ---
    _add_photo_centered(
        doc, portrait_doc.file_path if portrait_doc else None,
        width_cm=3.2, caption="Small ID Photo",
    )

    # --- application info chips ---
    _info_table(doc, [
        ("Application No", "رقم الطلب", f"#{profile.application_no}"),
        ("Post Applied For", "الوظيفة", profile.post_applied_for),
        ("Monthly Salary", "الراتب الشهري", f"{profile.monthly_salary} {profile.salary_currency}" if profile.monthly_salary else None),
        ("Contract Period", "مدة العقد", f"{profile.contract_period_years} yrs" if profile.contract_period_years else None),
    ])

    if profile.partner:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"Presented in partnership with: {profile.partner.name} — {profile.partner.country}")
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = INK_MUTED

    # --- name ---
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_before = Pt(14)
    name_run = name_p.add_run(applicant.full_name.upper())
    name_run.font.size = Pt(15)
    name_run.font.bold = True
    name_run.font.color.rgb = AUTHORITY

    # --- the 2 large, uncropped photos (Full-Length + Passport), same position as on the print page ---
    _add_photo_pair(
        doc,
        left=(full_doc.file_path if full_doc else None, "Full-Length Photo"),
        right=(passport_doc.file_path if passport_doc else None, "Passport Copy"),
    )

    # --- personal info ---
    _section_title(doc, "Personal Information", "المعلومات الشخصية")
    _info_table(doc, [
        ("Nationality", "الجنسية", "Ethiopian"),
        ("Religion", "الديانة", profile.religion),
        ("Date of Birth", "تاريخ الميلاد", applicant.date_of_birth.strftime("%d %m %Y") if applicant.date_of_birth else None),
        ("Place of Birth", "مكان الميلاد", profile.place_of_birth),
        ("Age", "العمر", profile.age()),
        ("Marital Status", "الحالة الاجتماعية", profile.marital_status),
        ("Weight / Height", "الوزن / الطول",
         f"{profile.weight_kg or '—'} kg / {profile.height_m or '—'} m"),
        ("Educational Qualification", "المؤهل العلمي", applicant.education_level),
    ])

    # --- passport ---
    _section_title(doc, "Passport Details", "بيانات جواز السفر")
    _info_table(doc, [
        ("Passport Number", "رقم الجواز", applicant.passport_number),
        ("Issue Place", "مكان الاصدار", profile.passport_issue_place),
        ("Issue Date", "تاريخ الاصدار", profile.passport_issue_date.strftime("%d/%m/%Y") if profile.passport_issue_date else None),
        ("Expiry Date", "تاريخ الانتهاء", profile.passport_expiry_date.strftime("%d/%m/%Y") if profile.passport_expiry_date else None),
    ])

    # --- languages ---
    _section_title(doc, "Languages", "اللغات")
    lang_pct = {"none": 10, "fair": 40, "good": 70, "fluent": 100}
    lang_map = {l["language"]: l["level"] for l in (profile.languages or [])}
    for lang in CVProfile.DEFAULT_LANGUAGES:
        _rating_row(doc, lang, lang_map.get(lang, "none"), lang_pct)

    # --- skills ---
    _section_title(doc, "Skills", "المهارات")
    skill_pct = {"poor": 25, "fair": 50, "good": 75, "excellent": 100}
    skill_map = {s["skill"]: s["level"] for s in (profile.skills or [])}
    for skill in CVProfile.DEFAULT_SKILLS:
        _rating_row(doc, skill, skill_map.get(skill, "good"), skill_pct)

    # --- work history ---
    if profile.work_history:
        _section_title(doc, "Work Experience", "الخبرة")
        table = doc.add_table(rows=1 + len(profile.work_history), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_cell_text(table.rows[0].cells[0], "Period", bold=True, size=9.5)
        _set_cell_text(table.rows[0].cells[1], "Country", bold=True, size=9.5)
        _shade_cell(table.rows[0].cells[0], CANVAS_GREY)
        _shade_cell(table.rows[0].cells[1], CANVAS_GREY)
        for i, h in enumerate(profile.work_history, start=1):
            _set_cell_text(table.rows[i].cells[0], h.get("period"), size=9.5)
            _set_cell_text(table.rows[i].cells[1], h.get("country"), size=9.5)
        _set_table_borders(table)

    # --- emergency contact ---
    _section_title(doc, "Emergency Contact", "معلومات الطوارئ")
    _info_table(doc, [
        ("Full Name", "الاسم الكامل", profile.emergency_contact_name),
        ("Address", "العنوان", profile.emergency_contact_address),
        ("Telephone", "رقم الهاتف", profile.emergency_contact_phone),
    ])

    # --- footer ---
    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(18)
    foot_run = foot.add_run(
        f"{company_name}  ·  {company_phone}  ·  {company_email}  ·  {company_address_en}"
    )
    foot_run.font.size = Pt(8)
    foot_run.font.color.rgb = INK_MUTED
    gen = doc.add_paragraph()
    gen_run = gen.add_run(
        f"Generated {profile.updated_at.strftime('%Y-%m-%d')}  ·  Application No. {profile.application_no}"
    )
    gen_run.font.size = Pt(8)
    gen_run.font.color.rgb = INK_MUTED

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
